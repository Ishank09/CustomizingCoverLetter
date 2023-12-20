from tkinter import *
import docx
import webbrowser
from tkinter import font
from docx.shared import Pt


# initialize window
root = Tk()
 
# set window title
root.title("Cover Letter Generator")
 
# set window size
root.geometry("1000x400")
 
# Create Labels
aspiration = Label(root, text = "Aspiration:")
role = Label(root, text = "Role:")
company = Label(root, text = "Company:")
top_skill = Label(root, text = "Top Skill:")
database = Label(root, text = "Databases:")
skills = Label(root, text = "Skills:")
concept = Label(root, text = "Concept:")
industry = Label(root, text = "Industry:")
tooltip = Label(root, text = "*If placeholder exceed 30 character, it will be given as a tool tip")

# Create textboxes
aspiration_txt = Entry(root, width=30)
role_txt = Entry(root, width=30)
company_txt = Entry(root, width=30)
top_skill_txt = Entry(root, width=30)
database_txt = Entry(root, width=30)
skills_txt = Entry(root, width=30)
concept_txt = Entry(root, width=30)
industry_txt = Entry(root, width=30)

# Set placeholder
aspiration_txt.insert(0, "Aspiring Software Engineer and Full Stack-Backend Developer")
role_txt.insert(0, "Software Engineer Intern")
company_txt.insert(0, "your esteemed organization")
top_skill_txt.insert(0, "Object-Oriented Programming, Restful APIs, microservices, databases, system design, and architecture")
database_txt.insert(0, "MS SQL, MySQL, MongoDB, and Elasticsearch")
skills_txt.insert(0, "Java, Python, C#, .Net Core, and Spring")
concept_txt.insert(0, "Object-Oriented Programming, Data Structures and Algorithms")
industry_txt.insert(0, "Software Engineering")

# create a button
generate_btn = Button(root, text = "Generate Cover Letter", command = lambda: generate_cover_letter(aspiration_txt.get(), role_txt.get(), company_txt.get(), top_skill_txt.get(), database_txt.get(), skills_txt.get(), concept_txt.get(), industry_txt.get()))

# Place the widget at respective positions
aspiration.grid(row=0, column=0)
role.grid(row=1, column=0)
company.grid(row=2, column=0)
top_skill.grid(row=3, column=0)
database.grid(row=4, column=0)
skills.grid(row=5, column=0)
concept.grid(row=6, column=0)
industry.grid(row=7, column=0)
tooltip.grid(row=8, column=0)

aspiration_txt.grid(row=0, column=1)
role_txt.grid(row=1, column=1)
company_txt.grid(row=2, column=1)
top_skill_txt.grid(row=3, column=1)
database_txt.grid(row=4, column=1)
skills_txt.grid(row=5, column=1)
concept_txt.grid(row=6, column=1)
industry_txt.grid(row=7, column=1)
generate_btn.grid(row=8, column=1)

def generate_cover_letter(aspiration, role, company, top_skill, database, skills, concept, industry):
    # open cv.docx
    doc = docx.Document('<PATH TO TEMLATE>/cv.docx')
    # font_name = font.Font(family='Times New Roman')

    # for each textarea in the form:
    # get the binding property from above, find it in the doc


    for para in doc.paragraphs:

        if '{Aspiring Software Engineer and Full Stack-Backend Developer}' in para.text:
            if aspiration == "{Aspiring Software Engineer and Full Stack-Backend Developer}":
                para.text = para.text.replace('{Aspiring Software Engineer and Full Stack-Backend Developer}', 'Aspiring Software Engineer and Full Stack-Backend Developer')
            else:
                para.text = para.text.replace('{Aspiring Software Engineer and Full Stack-Backend Developer}', aspiration)
        if '{ROLE}' in para.text:
                para.text = para.text.replace('{ROLE}', role)
        if '{COMPANY}' in para.text:
                para.text = para.text.replace('{COMPANY}', company)
        if '{TOP_SKILL}' in para.text:
                para.text = para.text.replace('{TOP_SKILL}', top_skill)
        if '{DB}' in para.text:
                para.text = para.text.replace('{DB}', database)
        if '{SKILL}' in para.text:
                para.text = para.text.replace('{SKILL}', skills)
        if '{CONCEPT}' in para.text:
                para.text = para.text.replace('{CONCEPT}', concept)
        if '{INDUSTRY}' in para.text:
                para.text = para.text.replace('{INDUSTRY}', industry)
    # save the resultin doc in the same location with name: {CPMPANY}_{ROLE}_CV


    doc.paragraphs[0].runs[0].font.size = docx.shared.Pt(22)
    doc.paragraphs[0].runs[0].font.name = 'Calibri'

    doc.paragraphs[1].runs[0].font.size = docx.shared.Pt(14)
    doc.paragraphs[1].runs[0].font.name = 'Calibri'
    for paragraph in doc.paragraphs[2:]:
        for run in paragraph.runs:
                run.font.size = docx.shared.Pt(12)
                run.font.name = 'Calibri'

    filename = company.replace(' ','.') + '_' + role.replace(' ','.') + '_CV.docx'
    print(filename)
    doc.save('<PATH TO SAVE GENERATED COVER LETTER>'+filename)

    # after the successful operation give a hyperlink to open that filepath in finder
    link = "file:///" + filename
#     print("Your cover letter is saved as " + filename + ". Click here to open: " + link)

    # Create Labels
    status = Label(root, text = "Customized Cover Letter Generated with name:"+filename)
#     link = Label(root, text = "Click here to open: " + link, fg="blue", cursor="hand2")
#     link.bind('<Button-1>', lambda e: callback(link))
    
    # Place the widget at respective positions
    status.grid(row=9, column=0)
    link.grid(row=9, column=1)

# callback function
def callback(link):
    webbrowser.open_new(link)


# run the main loop
root.mainloop()
